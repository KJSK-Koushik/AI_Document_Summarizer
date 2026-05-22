# utils/file_utils.py
import io
from pathlib import Path
import pdfplumber
from docx import Document

def extract_text_from_pdf(path_or_bytes):
    # path_or_bytes: filepath or bytes
    if isinstance(path_or_bytes, (bytes, bytearray)):
        fp = io.BytesIO(path_or_bytes)
    else:
        fp = path_or_bytes
    text_parts = []
    try:
        with pdfplumber.open(fp) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
    except Exception:
        # fallback: return empty string
        return ""
    return "\n".join(text_parts)


def extract_text_from_docx(path_or_bytes):
    if isinstance(path_or_bytes, (bytes, bytearray)):
        fp = io.BytesIO(path_or_bytes)
        doc = Document(fp)
    else:
        doc = Document(path_or_bytes)
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n".join(paragraphs)


def extract_text_from_txt(path_or_bytes):
    if isinstance(path_or_bytes, (bytes, bytearray)):
        return path_or_bytes.decode("utf-8", errors="ignore")
    else:
        return Path(path_or_bytes).read_text(encoding="utf-8", errors="ignore")


def extract_text(file):
    """file is a `pathlib.Path` or file-like object returned by Gradio's upload.
    Return extracted text (str).
    """
    mimetype = None
    name = getattr(file, 'name', '')
    # Gradio returns tempfile-like object with .name containing path
    try:
        suffix = Path(name).suffix.lower()
    except Exception:
        suffix = ''
    if suffix in ['.pdf']:
        return extract_text_from_pdf(file)
    elif suffix in ['.docx']:
        return extract_text_from_docx(file)
    elif suffix in ['.txt', '.text']:
        return extract_text_from_txt(file)
    else:
        # try pdf first
        txt = extract_text_from_pdf(file)
        if txt and len(txt.strip())>0:
            return txt
        # try docx
        txt = extract_text_from_docx(file)
        if txt and len(txt.strip())>0:
            return txt
        # last resort: try reading as text
        try:
            return extract_text_from_txt(file)
        except Exception:
            return ""